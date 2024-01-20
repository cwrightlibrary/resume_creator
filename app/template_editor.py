from docx import Document
from os.path import abspath, dirname, join, realpath

block_basic_folder = abspath(join(dirname(__file__), "..", "document-templates", "blocky-basic"))
save_folder = join(block_basic_folder, "edited")

resume = Document(join(block_basic_folder, "blocky-basic-resume.docx"))
cover_letter = Document(join(block_basic_folder, "blocky-basic-cover-letter.docx"))

# RESUME
_A_NAME = resume.paragraphs[0]
_B1_PHONE = resume.paragraphs[4]
_B2_EMAIL = resume.paragraphs[5]
_C_ABOUT = resume.paragraphs[8]

_D1_JOB_COMPANY = resume.paragraphs[11]
_D1_JOB_TITLE = resume.paragraphs[12]
_D1_JOB_DESCRIPTION = resume.paragraphs[13]
_D2_JOB_COMPANY = resume.paragraphs[14]
_D2_JOB_TITLE = resume.paragraphs[15]
_D2_JOB_DESCRIPTION = resume.paragraphs[16]
_D3_JOB_COMPANY = resume.paragraphs[17]
_D3_JOB_TITLE = resume.paragraphs[18]
_D3_JOB_DESCRIPTION = resume.paragraphs[19]
_D4_JOB_COMPANY = resume.paragraphs[20]
_D4_JOB_TITLE = resume.paragraphs[21]
_D4_JOB_DESCRIPTION = resume.paragraphs[22]

_E1_UNIVERSITY = resume.paragraphs[24]
_E1_UNIVERSITY_DESCRIPTION = resume.paragraphs[25]
_E2_HIGH_SCHOOL = resume.paragraphs[26]
_E2_HIGH_SCHOOL_DESCRIPTION = resume.paragraphs[27]

_F1_CERTIFICATION_INFO = resume.paragraphs[29]
_F2_CERTIFICATION_INFO = resume.paragraphs[30]
_F3_CERTIFICATION_INFO = resume.paragraphs[31]

_G_PHONE_EMAIL = resume.paragraphs[34]

_Z_HEADER = resume.sections[0].header.paragraphs[0]

_RESUME_ELEMENTS = [_A_NAME, _B1_PHONE, _B2_EMAIL, _C_ABOUT, _D1_JOB_COMPANY, _D1_JOB_TITLE, _D1_JOB_DESCRIPTION, _D2_JOB_COMPANY, _D2_JOB_TITLE, _D2_JOB_DESCRIPTION, _D3_JOB_COMPANY, _D3_JOB_TITLE, _D3_JOB_DESCRIPTION, _D4_JOB_COMPANY, _D4_JOB_TITLE, _D4_JOB_DESCRIPTION, _E1_UNIVERSITY, _E1_UNIVERSITY_DESCRIPTION, _E2_HIGH_SCHOOL, _E2_HIGH_SCHOOL_DESCRIPTION, _F1_CERTIFICATION_INFO, _F2_CERTIFICATION_INFO, _F3_CERTIFICATION_INFO, _G_PHONE_EMAIL, _Z_HEADER]

# COVER LETTER
_AA_NAME = cover_letter.paragraphs[0]
_BB_PHONE = cover_letter.paragraphs[4]
_CC_EMAIL = cover_letter.paragraphs[5]

_DD_HIRING_MANAGER_COMPANY_NAME_ADDRESS = cover_letter.paragraphs[11]
_EE_FROM_NAME_ADDRESS = cover_letter.paragraphs[13]

_FF1_TEXT = cover_letter.paragraphs[14]
_FF2_TEXT = cover_letter.paragraphs[15]
_FF3_TEXT = cover_letter.paragraphs[16]
_FF4_TEXT = cover_letter.paragraphs[17]

_ZZ_HEADER = cover_letter.sections[0].header.paragraphs[0]

_COVER_LETTER_ELEMENTS = [_AA_NAME, _BB_PHONE, _CC_EMAIL, _DD_HIRING_MANAGER_COMPANY_NAME_ADDRESS, _EE_FROM_NAME_ADDRESS, _FF1_TEXT, _FF2_TEXT, _FF3_TEXT, _FF4_TEXT, _ZZ_HEADER]

# For testing purposes
new_name = "Chris\nWright"
new_phone = "(302) 359-0237"
new_email = "jpchriswright@icloud.com"

_A_NAME.text = new_name
_B1_PHONE.text = new_phone
_B2_EMAIL.text = new_email
_G_PHONE_EMAIL.text = new_phone + "\n" + new_email
_Z_HEADER.text = _Z_HEADER.text.replace("First Surname", new_name.replace("\n", " "))

_AA_NAME.text = new_name
_BB_PHONE.text = new_phone
_CC_EMAIL.text = new_email
_EE_FROM_NAME_ADDRESS.text = new_name.replace("\n", " ")
_ZZ_HEADER.text = _ZZ_HEADER.text.replace("First Surname", new_name.replace("\n", " "))

# resume.save(join(save_folder, "resume-updated.docx"))
# cover_letter.save(join(save_folder, "cover-letter-udpated.docx"))
